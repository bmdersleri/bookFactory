# -*- coding: utf-8 -*-
"""
DOCX Güvenli Biçim Düzeltici
----------------------------

Bu betik, Pandoc ile üretilmiş DOCX dosyasında yalnızca hedeflenen
paragrafları biçimlendirir:

1. Resim içeren paragrafları ortalar.
2. Yalnızca ana bölüm/ek H1 başlıklarını ortalar.
3. Yalnızca standalone BÖLÜM SONU / EK ... SONU paragraflarını ortalar.
4. Dikkat, İpucu, Sınav Notu, Derinlemesine, Alıştırma Molası gibi
   pedagogik kutuları iki yana yaslar.
5. Tablo ilk satırlarını okunur hâle getirir:
   - açık gri arka plan
   - siyah yazı
   - kalın başlık
   - dikey ortalama
   - tekrar eden başlık satırı

Kullanım:
    python docx_guvenli_bicim_duzelt.py KompaktBirlesik.docx

Kopya çıktı üretmek için:
    python docx_guvenli_bicim_duzelt.py KompaktBirlesik.docx KompaktBirlesik_duzeltilmis.docx

Yerinde düzeltmek için:
    python docx_guvenli_bicim_duzelt.py KompaktBirlesik.docx --in-place

Tablo başlıklarına dokunmadan çalıştırmak için:
    python docx_guvenli_bicim_duzelt.py KompaktBirlesik.docx --no-table-header-fix
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


# ------------------------------------------------------------
# Temel yardımcılar
# ------------------------------------------------------------

def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def paragrafta_resim_var_mi(paragraph) -> bool:
    """
    Paragraf içinde Word resmi/drawing/pict var mı kontrol eder.
    """
    xml = paragraph._p.xml
    return (
        "<w:drawing" in xml
        or "<w:pict" in xml
        or "<pic:pic" in xml
        or "<a:blip" in xml
    )


def heading_1_mi(paragraph) -> bool:
    """
    Pandoc genellikle H1 başlıklarını Heading 1 / Başlık 1 stiline dönüştürür.
    Ek güvenlik olarak outlineLvl=0 da kontrol edilir.
    """
    style_name = paragraph.style.name if paragraph.style else ""

    if style_name in ("Heading 1", "Başlık 1"):
        return True

    xml = paragraph._p.xml
    if '<w:outlineLvl w:val="0"' in xml:
        return True

    return False


def ana_bolum_veya_ek_basligi_mi(paragraph) -> bool:
    """
    Sadece ana bölüm/ek başlıklarını yakalar.
    Örnek:
    Bölüm 1: ...
    Bölüm 01: ...
    Ek A: ...
    """
    text = paragraph_text(paragraph)
    if not text:
        return False

    if re.match(r"^(Bölüm|Bolum)\s+0?\d+\s*:", text, flags=re.IGNORECASE):
        return True

    if re.match(r"^Ek\s+[A-ZÇĞİÖŞÜ]\s*:", text, flags=re.IGNORECASE):
        return True

    return False


def bolum_sonu_mu(paragraph) -> bool:
    """
    Yalnızca tek başına duran BÖLÜM SONU veya EK ... SONU satırlarını yakalar.
    """
    text = paragraph_text(paragraph)
    text = text.replace("*", "").strip()

    if text == "BÖLÜM SONU":
        return True

    if re.match(r"^EK\s+[A-ZÇĞİÖŞÜ]\s+SONU$", text, flags=re.IGNORECASE):
        return True

    return False


def pedagogik_kutu_mu(paragraph) -> bool:
    """
    Dikkat, İpucu, Sınav Notu vb. kutu paragraflarını yakalar.
    Bunlar ortalanmamalı; iki yana yaslanmalıdır.
    """
    text = paragraph_text(paragraph)
    if not text:
        return False

    patterns = [
        r"^⚠️\s*Dikkat\s*:",
        r"^💡\s*İpucu\s*:",
        r"^💡\s*Ipucu\s*:",
        r"^🎯\s*Sınav Notu\s*:",
        r"^🎯\s*Sinav Notu\s*:",
        r"^🎯\s*Bölüm Hedefi\s*:",
        r"^🎯\s*Bolum Hedefi\s*:",
        r"^🔍\s*Derinlemesine\s*:",
        r"^Alıştırma Molası\s*:",
        r"^Alistirma Molasi\s*:",
        r"^Laboratuvar İpucu\s*:",
        r"^Laboratuvar Ipucu\s*:",
    ]

    return any(re.match(p, text, flags=re.IGNORECASE) for p in patterns)


def paragrafi_ortala(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def paragrafi_iki_yana_yasla(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def paragrafi_sola_yasla(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def iter_all_tables(doc):
    """
    Ana gövdedeki ve iç içe tablolardaki tüm tabloları dolaşır.
    """
    def nested_tables(table):
        yield table
        for row in table.rows:
            for cell in row.cells:
                for t in cell.tables:
                    yield from nested_tables(t)

    for table in doc.tables:
        yield from nested_tables(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


# ------------------------------------------------------------
# Tablo başlığı görünürlük düzeltmesi
# ------------------------------------------------------------

def set_cell_shading(cell, fill: str = "D9EAF7") -> None:
    """
    Hücre arka plan rengini ayarlar.
    fill örn: D9EAF7 açık mavi-gri.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)


def mark_row_as_repeating_header(row) -> None:
    """
    Word tablolarında ilk satırı tekrar eden başlık satırı olarak işaretler.
    """
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))

    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)

    tbl_header.set(qn("w:val"), "true")


def fix_table_header_row(table) -> int:
    """
    İlk satırı okunur tablo başlığı hâline getirir.
    Döndürdüğü değer: düzenlenen hücre sayısı.
    """
    if not table.rows:
        return 0

    header_row = table.rows[0]
    mark_row_as_repeating_header(header_row)

    count = 0

    for cell in header_row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")

        for p in cell.paragraphs:
            # Tablo başlıkları merkezli olabilir; bu pedagogik kutu değildir.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        count += 1

    return count


# ------------------------------------------------------------
# Ana düzeltme fonksiyonu
# ------------------------------------------------------------

def docx_guvenli_bicim_duzelt(
    input_docx: Path,
    output_docx: Path,
    fix_table_headers: bool = True,
) -> dict:
    doc = Document(str(input_docx))

    stats = {
        "centered_images": 0,
        "centered_h1": 0,
        "centered_end_markers": 0,
        "justified_pedagogic_boxes": 0,
        "fixed_table_header_cells": 0,
        "left_restored_non_target_centered_boxes": 0,
    }

    # Ana gövdedeki paragraflar.
    # Burada tablo içi paragraflara dokunmuyoruz.
    for p in doc.paragraphs:
        if pedagogik_kutu_mu(p):
            paragrafi_iki_yana_yasla(p)
            stats["justified_pedagogic_boxes"] += 1
            continue

        if paragrafta_resim_var_mi(p):
            paragrafi_ortala(p)
            stats["centered_images"] += 1
            continue

        if heading_1_mi(p) and ana_bolum_veya_ek_basligi_mi(p):
            paragrafi_ortala(p)
            stats["centered_h1"] += 1
            continue

        if bolum_sonu_mu(p):
            paragrafi_ortala(p)
            stats["centered_end_markers"] += 1
            continue

        # Önceki hatalı betik yüzünden kutu dışındaki bazı özel paragraflar ortalanmışsa,
        # sadece pedagogik kutulara dokunduğumuz için burada genel bir sola alma yapmıyoruz.
        # Genel sola alma, kitapta bilinçli ortalanmış paragrafları bozabilir.

    # Tablo içi paragraflar.
    # H1 veya pedagogik kutu aramayız. Tablo içi metinlerin hizalamasını genel olarak bozmayız.
    for table in iter_all_tables(doc):
        if fix_table_headers:
            stats["fixed_table_header_cells"] += fix_table_header_row(table)

        # Tablo içinde resim varsa yalnızca resim bulunan paragraf ortalanır.
        # Diğer tablo hücreleri değiştirilmez.
        for p in iter_table_paragraphs(table):
            if paragrafta_resim_var_mi(p):
                paragrafi_ortala(p)
                stats["centered_images"] += 1

    doc.save(str(output_docx))
    return stats


# ------------------------------------------------------------
# Komut satırı
# ------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Pandoc DOCX dosyasında resim, ana başlık, bölüm sonu, pedagogik kutu ve tablo başlığı biçimlerini güvenli düzeltir."
    )
    parser.add_argument("input_docx", help="Girdi DOCX dosyası")
    parser.add_argument("output_docx", nargs="?", help="Çıktı DOCX dosyası")
    parser.add_argument("--in-place", action="store_true", help="Aynı dosyanın üzerine yazar; önce .bak yedeği alır")
    parser.add_argument("--no-table-header-fix", action="store_true", help="Tablo başlık satırlarını düzeltmez")
    args = parser.parse_args()

    input_docx = Path(args.input_docx)

    if not input_docx.exists():
        print(f"Hata: Dosya bulunamadı: {input_docx}")
        return 1

    if args.in_place:
        backup = input_docx.with_suffix(input_docx.suffix + ".bak")
        shutil.copy2(input_docx, backup)
        output_docx = input_docx
        print(f"Yedek oluşturuldu: {backup}")
    else:
        output_docx = Path(args.output_docx) if args.output_docx else input_docx.with_name(input_docx.stem + "_duzeltilmis.docx")

    stats = docx_guvenli_bicim_duzelt(
        input_docx=input_docx,
        output_docx=output_docx,
        fix_table_headers=not args.no_table_header_fix,
    )

    print("DOCX güvenli biçim düzeltme tamamlandı.")
    print(f"Girdi : {input_docx}")
    print(f"Çıktı : {output_docx}")
    print()
    print("İşlem özeti:")
    for k, v in stats.items():
        print(f"  {k}: {v}")

    print()
    print("Not: Final teslimden önce DOCX'i Word/LibreOffice ile açıp tablo başlıklarını ve kutu hizalarını görsel olarak kontrol edin.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
